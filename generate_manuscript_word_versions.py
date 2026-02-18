#!/usr/bin/env python3
"""
Generate two versions of manuscript in Word format:
1. Clean Version - for publication
2. Tracked Changes Version - with highlighted sections for revision
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def add_highlight(paragraph, color='yellow'):
    """Add highlight to paragraph"""
    for run in paragraph.runs:
        run.font.highlight_color = 3 if color == 'yellow' else 4  # 3=yellow, 4=green

def create_manuscript_content(doc, tracked_changes=False):
    """Create manuscript content, optionally with tracked changes highlighting"""
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Deep Learning-Based Earthquake Precursor Detection from Geomagnetic Data: A Comparative Study of VGG16 and EfficientNet Architectures")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors
    authors = doc.add_paragraph()
    authors.add_run("Sumawan").bold = True
    authors.add_run("¹,², Bambang L. Widjiantoro¹, Katherin Indriawati¹, Muhamad Syirojudin²")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Affiliations
    aff = doc.add_paragraph()
    aff.add_run("¹Department of Engineering Physics, Sepuluh Nopember Institute of Technology, Surabaya 60111, Indonesia\n")
    aff.add_run("²Meteorological, Climatological and Geophysical Agency (BMKG), Jakarta, Indonesia")
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_run = aff.runs[0]
    aff_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract_text = """Earthquake prediction remains a grand challenge in geophysics. This study presents a comprehensive deep learning approach for automated earthquake precursor detection from geomagnetic spectrogram data, comparing two state-of-the-art convolutional neural network architectures—VGG16 and EfficientNet-B0—for multi-task learning of earthquake magnitude (4 classes) and azimuth direction (9 classes) prediction. Our dataset comprises 256 unique earthquake events (M4.0–7.0+) from 25 geomagnetic stations across Indonesia (2018–2025), generating 1,972 spectrogram samples through temporal windowing. Using fixed data splitting to prevent leakage, VGG16 achieved 98.68% magnitude and 54.93% azimuth accuracy, while EfficientNet-B0 achieved 94.37% magnitude and 57.39% azimuth accuracy with 26× fewer parameters (20 MB vs 528 MB) and 2.5× faster inference. Leave-One-Event-Out (LOEO) cross-validation confirmed robust generalization with acceptable performance drops (<5%). Gradient-weighted Class Activation Mapping (Grad-CAM) visualizations revealed that both models focus on ultra-low frequency (ULF) bands (0.001–0.01 Hz), consistent with geomagnetic precursor theory, demonstrating physically meaningful feature learning."""
    abstract_para = doc.add_paragraph(abstract_text)
    
    # Keywords
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("Earthquake precursor, deep learning, geomagnetic data, VGG16, EfficientNet, multi-task learning, Grad-CAM, LOEO validation")
    
    doc.add_paragraph()
    
    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    intro_text = """Earthquake prediction remains one of the most challenging problems in geophysics, with significant implications for disaster preparedness and risk mitigation. Among various precursor signals, Ultra-Low Frequency (ULF) geomagnetic anomalies have shown promising correlations with seismic activity. These signals, typically in the 0.001–1 Hz frequency range, are believed to originate from stress-induced electromagnetic emissions in the Earth's crust prior to major earthquakes.

Recent advances in deep learning have opened new possibilities for automated detection and classification of earthquake precursors. Convolutional Neural Networks (CNNs) have demonstrated remarkable success in image classification tasks, making them suitable candidates for analyzing spectrogram representations of geomagnetic signals.

This study addresses this gap by comparing two representative CNN architectures: VGG16, representing classical deep CNN design, and EfficientNet-B0, representing modern efficient architecture design. Our contributions include:

1. First comprehensive comparison of VGG16 and EfficientNet-B0 for geomagnetic earthquake precursor detection
2. Multi-task learning framework for simultaneous magnitude and azimuth prediction
3. Rigorous Leave-One-Event-Out (LOEO) validation ensuring generalization to unseen earthquake events
4. Explainable AI analysis using Grad-CAM demonstrating physically meaningful feature learning"""
    doc.add_paragraph(intro_text)
    
    # 2. Related Work
    doc.add_heading("2. Related Work", level=1)
    related_text = """The relationship between geomagnetic anomalies and seismic activity has been studied extensively. Hayakawa et al. demonstrated that ULF emissions in the 0.001–0.01 Hz range show enhanced activity before major earthquakes. Hattori provided comprehensive evidence for ULF geomagnetic changes associated with large earthquakes in Japan.

Machine learning approaches for earthquake-related tasks have gained significant attention. Han et al. applied LSTM networks for earthquake prediction using geomagnetic data. Akhoondzadeh employed CNN-based methods for detecting ionospheric precursors. However, systematic comparisons between different architectures remain scarce."""
    doc.add_paragraph(related_text)
    
    # 3. Dataset and Preprocessing
    doc.add_heading("3. Dataset and Preprocessing", level=1)
    dataset_text = """Our dataset comprises ULF geomagnetic recordings from 25 stations operated by BMKG (Badan Meteorologi, Klimatologi, dan Geofisika) Indonesia. The three-component magnetometer data (H, D, Z) was collected at 1 Hz sampling rate.

We selected 256 unique earthquake events with magnitude M4.0–7.0+ occurring within 500 km of the geomagnetic stations during 2018–2025. The dataset includes:
- Moderate (M4.0–4.9): 20 events
- Medium (M5.0–5.9): 1,036 samples
- Large (M6.0+): 28 samples
- Normal (No earthquake): 888 samples
- Total: 1,972 spectrogram samples

Raw geomagnetic data was processed through bandpass filtering (0.001–0.5 Hz), temporal segmentation (6-hour windows), STFT spectrogram generation, and normalization using ImageNet statistics."""
    doc.add_paragraph(dataset_text)
    
    # 4. Methodology
    doc.add_heading("4. Methodology", level=1)
    method_text = """We formulate earthquake precursor detection as a multi-task classification problem where the input spectrogram is mapped to both magnitude class (4 classes) and azimuth direction class (9 classes).

VGG16 consists of 16 weight layers with 3×3 convolution filters, modified for multi-task learning with separate heads for magnitude and azimuth prediction. Total parameters: 138M (528 MB model size).

EfficientNet-B0 uses compound scaling to balance network depth, width, and resolution. The multi-task head uses FC(1280→512→4/9) with dropout 0.444. Total parameters: 5.3M (20 MB model size).

Both models were trained with Adam optimizer, Focal Loss (γ=2) for class imbalance handling, batch size 32, and early stopping with 10 epochs patience. Data split: 70/15/15 (train/validation/test) with fixed seed."""
    doc.add_paragraph(method_text)
    
    # 5. Results
    doc.add_heading("5. Experimental Results", level=1)
    results_text = """Performance Comparison:
- VGG16: 98.68% magnitude accuracy, 54.93% azimuth accuracy, 528 MB, 125 ms inference
- EfficientNet-B0: 94.37% magnitude accuracy, 57.39% azimuth accuracy, 20 MB, 50 ms inference

VGG16 achieves higher magnitude accuracy (+4.31%), while EfficientNet-B0 shows better azimuth accuracy (+2.46%) with significantly smaller model size (26× reduction) and faster inference (2.5× speedup).

LOEO Cross-Validation Results:
- Random Split: 98.68% magnitude, 54.93% azimuth
- LOEO (10-fold): 94.23% ± 2.1% magnitude, 52.18% ± 3.4% azimuth
- Performance Drop: 4.45% (within acceptable <5% limit)

Grad-CAM Analysis confirmed that both models focus on ULF frequency bands (0.001–0.01 Hz), temporal evolution patterns, and magnitude-dependent signal intensity variations."""
    doc.add_paragraph(results_text)
    
    # 6. Discussion
    doc.add_heading("6. Discussion", level=1)
    discussion_text = """VGG16 achieves maximum accuracy (98.68%) but requires significant computational resources (528 MB, 125 ms). EfficientNet-B0 offers 94.37% accuracy with 26× smaller footprint, making it suitable for edge deployment and real-time monitoring applications.

Both models show lower azimuth accuracy (~55%) compared to magnitude (~97%). This reflects the inherent difficulty of 9-class directional classification from geomagnetic signals. However, 55% significantly exceeds the random baseline (11.1%), demonstrating meaningful directional learning.

Limitations include limited samples for rare classes, regional specificity (Indonesia), and azimuth accuracy requiring further improvement.

For production early warning systems, EfficientNet-B0 is recommended due to acceptable accuracy, small footprint, fast inference, and suitability for real-time monitoring."""
    doc.add_paragraph(discussion_text)
    
    # 7. Conclusion
    doc.add_heading("7. Conclusion", level=1)
    conclusion_text = """This study compared VGG16 and EfficientNet-B0 for earthquake precursor detection from geomagnetic spectrograms. Key findings:

1. VGG16 achieves 98.68% magnitude accuracy (best overall)
2. EfficientNet-B0 achieves 94.37% with 26× smaller model
3. LOEO validation confirms <5% generalization drop
4. Grad-CAM reveals physically meaningful ULF band focus

EfficientNet-B0 offers optimal balance for operational early warning systems, enabling real-time precursor detection with minimal computational resources. Future work will expand the dataset with more Large events and improve azimuth classification accuracy."""
    doc.add_paragraph(conclusion_text)
    
    # Data Availability (NEW SECTION - highlight if tracked changes)
    doc.add_heading("Data Availability Statement", level=1)
    data_avail = """The data and computational resources supporting the findings of this study are handled as follows:

Geomagnetic Data: Raw three-component (H, D, Z) geomagnetic time-series data are provided by the Agency for Meteorology, Climatology, and Geophysics (BMKG), Indonesia. Access to raw data is subject to institutional data-sharing policies.

Earthquake Catalog: Seismic event metadata were obtained from the BMKG seismic bulletin and the USGS Earthquake Hazards Program.

Processed Dataset: A curated subset of 1,972 normalized STFT spectrogram samples is available for academic replication.

Repository Access: The source code has been made publicly available at: https://github.com/sumawanbmkg/earthquake-precursor-cnn"""
    
    data_para = doc.add_paragraph(data_avail)
    if tracked_changes:
        # Highlight this section as NEW
        for run in data_para.runs:
            run.font.highlight_color = 7  # Yellow highlight
    
    # Acknowledgments
    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph("The authors thank BMKG (Badan Meteorologi, Klimatologi, dan Geofisika) Indonesia for providing geomagnetic and earthquake catalog data.")
    
    # References
    doc.add_heading("References", level=1)
    refs = """[1] M. Hayakawa et al., "Criticality features in ULF magnetic fields prior to the 2011 Tohoku earthquake," Proc. Jpn. Acad., Ser. B, vol. 91, no. 1, pp. 25–30, 2015.
[2] K. Hattori, "ULF geomagnetic changes associated with large earthquakes," Terr. Atmos. Ocean. Sci., vol. 15, no. 3, pp. 329–360, 2004.
[3] K. Simonyan and A. Zisserman, "Very deep convolutional networks for large-scale image recognition," in Proc. ICLR, 2015.
[4] M. Tan and Q. Le, "EfficientNet: Rethinking model scaling for convolutional neural networks," in Proc. ICML, 2019, pp. 6105–6114.
[5] T.-Y. Lin et al., "Focal loss for dense object detection," in Proc. IEEE ICCV, 2017, pp. 2980–2988.
[6] R. R. Selvaraju et al., "Grad-CAM: Visual explanations from deep networks," in Proc. IEEE ICCV, 2017, pp. 618–626."""
    doc.add_paragraph(refs)
    
    return doc

# Create output directory
output_dir = Path('publication/paper')
output_dir.mkdir(parents=True, exist_ok=True)

# 1. Generate Clean Version
doc_clean = Document()
create_manuscript_content(doc_clean, tracked_changes=False)
clean_path = output_dir / 'Manuscript_Clean_Version.docx'
doc_clean.save(clean_path)

# 2. Generate Tracked Changes Version (with highlights)
doc_tracked = Document()
create_manuscript_content(doc_tracked, tracked_changes=True)

# Add note at the beginning for tracked changes version
first_para = doc_tracked.paragraphs[0]
note = doc_tracked.add_paragraph()
note_run = note.add_run("[NOTE: Yellow highlighted sections indicate NEW or REVISED content added in this revision. This file is for peer review only and will not be published.]")
note_run.font.color.rgb = RGBColor(255, 0, 0)
note_run.font.size = Pt(10)
note_run.italic = True
# Move note to beginning
note._p.addprevious(first_para._p)

tracked_path = output_dir / 'Manuscript_Tracked_Changes.docx'
doc_tracked.save(tracked_path)

# Copy to publication_package
import shutil
pkg_dir = Path('publication_package')
pkg_dir.mkdir(exist_ok=True)
shutil.copy(clean_path, pkg_dir / 'Manuscript_Clean_Version.docx')
shutil.copy(tracked_path, pkg_dir / 'Manuscript_Tracked_Changes.docx')

print("=" * 70)
print("Manuscript Word versions generated successfully!")
print("=" * 70)
print(f"\nOutput files:")
print(f"\n1. CLEAN VERSION (for publication):")
print(f"   - {clean_path}")
print(f"   - {pkg_dir / 'Manuscript_Clean_Version.docx'}")
print(f"\n2. TRACKED CHANGES VERSION (for peer review):")
print(f"   - {tracked_path}")
print(f"   - {pkg_dir / 'Manuscript_Tracked_Changes.docx'}")
print("\n" + "=" * 70)
print("USAGE NOTES:")
print("=" * 70)
print("""
For INITIAL SUBMISSION:
  → Upload only the Clean Version

For REVISION SUBMISSION:
  → Upload Clean Version (main manuscript)
  → Upload Tracked Changes Version (shows revisions)
  
The Tracked Changes version has:
  - Yellow highlights on NEW sections (e.g., Data Availability Statement)
  - Red note at top indicating it's for review only
  
To add more tracked changes later:
  1. Open in Microsoft Word
  2. Enable Review > Track Changes
  3. Make your edits (they will be tracked automatically)
""")
