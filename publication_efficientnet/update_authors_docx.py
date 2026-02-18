#!/usr/bin/env python3
"""
Update All DOCX Files with Real Author Information

Author: Research Team
Date: February 14, 2026
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Author Information
AUTHORS = [
    {
        'order': '1st',
        'first_name': 'Sumawan',
        'last_name': 'Sumawan',
        'email': 'sumawanbmkg@gmail.com',
        'phone': '+62 085239395272',
        'country': 'Indonesia',
        'affiliation': 'Sepuluh Nopember Institute of Technology, Surabaya, Indonesia',
        'orcid': 'https://orcid.org/0009-0005-5301-6414',
        'role': 'Corresponding Author'
    },
    {
        'order': '2nd',
        'first_name': 'Bambang L.',
        'last_name': 'Widjiantoro',
        'email': 'bambang.lw@its.ac.id',
        'phone': '+62 08179670328',
        'country': 'Indonesia',
        'affiliation': 'Sepuluh Nopember Institute of Technology, Surabaya, Indonesia',
        'orcid': 'https://orcid.org/0009-0003-1000-3184',
        'role': 'Co-Author'
    },
    {
        'order': '3rd',
        'first_name': 'Katherin',
        'last_name': 'Indriawati',
        'email': 'katherin@ep.its.ac.id',
        'phone': '+62 085258747305',
        'country': 'Indonesia',
        'affiliation': 'Sepuluh Nopember Institute of Technology, Surabaya, Indonesia',
        'orcid': 'https://orcid.org/0000-0002-9333-088X',
        'role': 'Co-Author'
    },
    {
        'order': '4th',
        'first_name': 'Muhamad',
        'last_name': 'Syirojudin',
        'email': 'muhamad.syirojudin@bmkg.go.id',
        'phone': '+62 081382051913',
        'country': 'Indonesia',
        'affiliation': 'Meteorological, Climatological and Geophysical Agency, Jakarta, Indonesia',
        'orcid': 'https://orcid.org/0000-0002-3170-7223',
        'role': 'Co-Author'
    }
]

def get_author_full_name(author):
    """Get full name of author"""
    return f"{author['first_name']} {author['last_name']}"

def get_author_list_text():
    """Get formatted author list"""
    return ', '.join([get_author_full_name(a) for a in AUTHORS])

def get_corresponding_author():
    """Get corresponding author"""
    return AUTHORS[0]

def update_declaration_docx():
    """Update Declaration of Interest with real authors"""
    try:
        doc = Document('1_DECLARATION_OF_INTEREST.docx')
        
        # Update title
        doc.paragraphs[0].text = 'Declaration of Competing Interest'
        
        # Update manuscript info
        doc.paragraphs[1].text = (
            'Manuscript Title: Hierarchical EfficientNet for Earthquake Precursor Detection '
            'from Geomagnetic Spectrograms: A Multi-Task Deep Learning Approach'
        )
        
        # Find and update authors section
        for i, para in enumerate(doc.paragraphs):
            if 'Authors:' in para.text:
                # Clear old author list
                for j in range(i+1, min(i+5, len(doc.paragraphs))):
                    if doc.paragraphs[j].style.name == 'List Bullet':
                        doc.paragraphs[j].text = ''
                
                # Add new authors
                for idx, author in enumerate(AUTHORS):
                    if i+1+idx < len(doc.paragraphs):
                        doc.paragraphs[i+1+idx].text = (
                            f"{get_author_full_name(author)}, {author['affiliation']}"
                        )
                break
        
        # Update corresponding author at end
        for para in doc.paragraphs:
            if 'Corresponding Author Name:' in para.text:
                para.text = f"Corresponding Author Name: {get_author_full_name(AUTHORS[0])}"
            elif 'Corresponding Author Email:' in para.text:
                para.text = f"Corresponding Author Email: {AUTHORS[0]['email']}"
        
        doc.save('1_DECLARATION_OF_INTEREST.docx')
        print("✅ Updated: 1_DECLARATION_OF_INTEREST.docx")
        
    except Exception as e:
        print(f"❌ Error updating Declaration: {e}")

def update_cover_letter_docx():
    """Update Cover Letter with real authors"""
    try:
        doc = Document('3_COVER_LETTER.docx')
        
        # Update corresponding author section at end
        updated = False
        for i, para in enumerate(doc.paragraphs):
            if 'Corresponding Author Name' in para.text or '[Corresponding Author Name]' in para.text:
                doc.paragraphs[i].text = get_author_full_name(AUTHORS[0])
                updated = True
            elif '[Title/Position]' in para.text:
                doc.paragraphs[i].text = 'Ph.D. Student / Researcher'
            elif 'Email: [email@its.ac.id]' in para.text:
                doc.paragraphs[i].text = f"Email: {AUTHORS[0]['email']}"
            elif 'Co-Authors' in para.text:
                # Update co-authors list
                if i+1 < len(doc.paragraphs):
                    co_authors_text = '\n'.join([
                        f"**{get_author_full_name(a)}**: {a['affiliation'].split(',')[0]}"
                        for a in AUTHORS[1:]
                    ])
                    doc.paragraphs[i+1].text = co_authors_text
        
        doc.save('3_COVER_LETTER.docx')
        print("✅ Updated: 3_COVER_LETTER.docx")
        
    except Exception as e:
        print(f"❌ Error updating Cover Letter: {e}")

def create_author_info_docx():
    """Create comprehensive author information document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Author Information', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'Hierarchical EfficientNet for Earthquake Precursor Detection',
        style='Subtitle'
    )
    
    doc.add_paragraph()
    
    # Author list
    doc.add_heading('Authors', level=1)
    
    for author in AUTHORS:
        doc.add_heading(f"{author['order']} Author: {get_author_full_name(author)}", level=2)
        
        # Create info table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info = [
            ('Name', get_author_full_name(author)),
            ('Email', author['email']),
            ('Phone', author['phone']),
            ('Affiliation', author['affiliation']),
            ('Country', author['country']),
            ('ORCID', author['orcid']),
            ('Role', author['role'])
        ]
        
        for i, (key, value) in enumerate(info):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
    
    # Corresponding author section
    doc.add_page_break()
    doc.add_heading('Corresponding Author Details', level=1)
    
    corr = AUTHORS[0]
    doc.add_paragraph(f"**Name**: {get_author_full_name(corr)}")
    doc.add_paragraph(f"**Email**: {corr['email']}")
    doc.add_paragraph(f"**Phone**: {corr['phone']}")
    doc.add_paragraph(f"**Affiliation**: {corr['affiliation']}")
    doc.add_paragraph(f"**ORCID**: {corr['orcid']}")
    
    doc.add_paragraph()
    
    # Author contributions
    doc.add_heading('Author Contributions', level=1)
    
    contributions = {
        'Sumawan Sumawan': 'Conceptualization, Methodology, Software, Validation, Formal Analysis, Investigation, Data Curation, Writing - Original Draft, Visualization',
        'Bambang L. Widjiantoro': 'Supervision, Project Administration, Writing - Review & Editing, Resources',
        'Katherin Indriawati': 'Supervision, Writing - Review & Editing, Methodology',
        'Muhamad Syirojudin': 'Data Curation, Resources, Validation, Writing - Review & Editing'
    }
    
    for author_name, contribution in contributions.items():
        doc.add_paragraph(f"**{author_name}**: {contribution}", style='List Bullet')
    
    doc.add_paragraph()
    
    # Affiliations
    doc.add_heading('Affiliations', level=1)
    
    affiliations = {
        '1': 'Sepuluh Nopember Institute of Technology (ITS), Surabaya, Indonesia',
        '2': 'Meteorological, Climatological and Geophysical Agency (BMKG), Jakarta, Indonesia'
    }
    
    for key, aff in affiliations.items():
        doc.add_paragraph(f"**{key}**: {aff}", style='List Number')
    
    # Save
    doc.save('0_AUTHOR_INFORMATION.docx')
    print("✅ Created: 0_AUTHOR_INFORMATION.docx")

def create_author_list_for_manuscript():
    """Create formatted author list for manuscript"""
    doc = Document()
    
    # Title
    doc.add_heading('Author List for Manuscript', 0)
    
    # Formatted for manuscript
    doc.add_heading('Format 1: Full Names with Affiliations', level=2)
    
    author_text = []
    for i, author in enumerate(AUTHORS, 1):
        superscript = '1' if 'ITS' in author['affiliation'] else '2'
        author_text.append(f"{get_author_full_name(author)}{superscript}")
    
    doc.add_paragraph(', '.join(author_text))
    
    doc.add_paragraph()
    doc.add_paragraph('¹Sepuluh Nopember Institute of Technology, Surabaya, Indonesia')
    doc.add_paragraph('²Meteorological, Climatological and Geophysical Agency, Jakarta, Indonesia')
    
    doc.add_paragraph()
    
    # Format 2: With ORCID
    doc.add_heading('Format 2: With ORCID IDs', level=2)
    
    for author in AUTHORS:
        doc.add_paragraph(
            f"{get_author_full_name(author)} ({author['orcid']})",
            style='List Bullet'
        )
    
    doc.add_paragraph()
    
    # Format 3: For submission system
    doc.add_heading('Format 3: For Online Submission System', level=2)
    
    table = doc.add_table(rows=len(AUTHORS)+1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ['Order', 'First Name', 'Last Name', 'Email', 'Affiliation', 'ORCID']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Data
    for i, author in enumerate(AUTHORS, 1):
        table.rows[i].cells[0].text = author['order']
        table.rows[i].cells[1].text = author['first_name']
        table.rows[i].cells[2].text = author['last_name']
        table.rows[i].cells[3].text = author['email']
        table.rows[i].cells[4].text = author['affiliation']
        table.rows[i].cells[5].text = author['orcid'].split('/')[-1]
    
    # Save
    doc.save('0_AUTHOR_LIST_FORMATS.docx')
    print("✅ Created: 0_AUTHOR_LIST_FORMATS.docx")

def main():
    """Update all DOCX files with author information"""
    print("=" * 60)
    print("Updating DOCX Files with Real Author Information")
    print("=" * 60)
    print()
    
    # Change to publication folder
    os.chdir(os.path.dirname(__file__))
    
    try:
        # Update existing files
        update_declaration_docx()
        update_cover_letter_docx()
        
        # Create new author info files
        create_author_info_docx()
        create_author_list_for_manuscript()
        
        print()
        print("=" * 60)
        print("✅ All files updated successfully!")
        print("=" * 60)
        print()
        print("Updated/Created files:")
        print("  1. 1_DECLARATION_OF_INTEREST.docx (updated)")
        print("  2. 3_COVER_LETTER.docx (updated)")
        print("  3. 0_AUTHOR_INFORMATION.docx (new)")
        print("  4. 0_AUTHOR_LIST_FORMATS.docx (new)")
        print()
        print("Author Information:")
        print(f"  Corresponding Author: {get_author_full_name(AUTHORS[0])}")
        print(f"  Email: {AUTHORS[0]['email']}")
        print(f"  Total Authors: {len(AUTHORS)}")
        print()
        print("Next steps:")
        print("  1. Review all updated files")
        print("  2. Copy author list to main manuscript")
        print("  3. Verify ORCID IDs are correct")
        print("  4. Check affiliations formatting")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
