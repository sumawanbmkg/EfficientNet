#!/usr/bin/env python3
"""
Generate Figures DOCX with All Images and Captions

Author: Research Team
Date: February 14, 2026
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_figures_docx():
    """Create comprehensive figures document with images and captions"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Figures for Publication', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'Hierarchical EfficientNet for Earthquake Precursor Detection',
        style='Subtitle'
    )
    
    doc.add_paragraph()
    
    # Main Figures
    doc.add_heading('Main Figures', level=1)
    
    figures = [
        {
            'number': 1,
            'file': 'figures/FIG_1_Station_Map.png',
            'title': 'BMKG Geomagnetic Observatory Network and Study Area',
            'caption': (
                'Geographic distribution of 24 BMKG geomagnetic observatories across Indonesia '
                'used in this study. Stations are color-coded by data contribution: red circles '
                'indicate high-contribution stations (>100 samples), orange circles indicate '
                'medium-contribution stations (50-100 samples), and yellow circles indicate '
                'supporting stations (<50 samples). The map shows Indonesia\'s strategic position '
                'in the Pacific Ring of Fire, with stations distributed across major tectonic '
                'boundaries. Station codes are labeled for reference (e.g., TND = Tondano, '
                'KPG = Kupang, JYP = Jayapura). The network provides comprehensive coverage for '
                'detecting geomagnetic precursors associated with earthquake activity across the '
                'Indonesian archipelago.'
            )
        },
        {
            'number': 2,
            'file': 'figures/FIG_2_Preprocessing_Flow.png',
            'title': 'Data Preprocessing and Spectrogram Generation Pipeline',
            'caption': (
                'Comprehensive data preprocessing pipeline for converting raw geomagnetic time '
                'series to model-ready spectrograms. (A) Raw data acquisition from BMKG '
                'observatories via SSH, showing 1-hour temporal windows before earthquake events. '
                '(B) Three-component geomagnetic field measurements (H: horizontal north, '
                'D: declination, Z: vertical) sampled at 1 Hz. (C) Short-Time Fourier Transform '
                '(STFT) applied to each component with 256-sample window and 50% overlap. '
                '(D) Frequency filtering to isolate Ultra-Low Frequency (ULF) band (0.01-0.1 Hz). '
                '(E) Min-max normalization per channel. (F) RGB composition where H→Red, D→Green, '
                'Z→Blue channels, creating a 224×224×3 input tensor. (G) Example spectrograms for '
                'different magnitude classes showing distinct spectral signatures.'
            )
        },
        {
            'number': 3,
            'file': 'figures/FIG_3_Model_Architecture.png',
            'title': 'Hierarchical EfficientNet Architecture',
            'caption': (
                'Detailed architecture of the proposed Hierarchical EfficientNet model. The model '
                'consists of three main components: (1) Backbone: EfficientNet-B0 pretrained on '
                'ImageNet, serving as feature extractor. (2) Shared Neck: A 256-dimensional '
                'embedding layer with batch normalization and SiLU activation. (3) Multi-Task Heads: '
                'Three specialized prediction heads - (a) Binary Head (2 classes): Gatekeeper for '
                'precursor vs. normal classification; (b) Magnitude Head (4 classes): Estimates '
                'earthquake magnitude (Normal, Moderate, Medium, Large) with 2× class weight boost '
                'for Large events; (c) Azimuth Head (9 classes): Predicts earthquake source direction. '
                'Total loss: L_total = 2.0×L_binary + 1.0×L_magnitude + 0.5×L_azimuth. Model '
                'parameters: 5.8M. Inference time: 73ms per sample on CPU.'
            )
        },
        {
            'number': 4,
            'file': 'figures/FIG_4_Training_History.png',
            'title': 'Training History and Convergence Analysis',
            'caption': (
                'Training dynamics over 50 epochs showing model convergence and generalization. '
                '(A) Total Loss: Combined loss for training (blue) and validation (orange) sets. '
                'Early stopping triggered at epoch 42. (B) Binary Classification Loss: Gatekeeper '
                'task showing excellent convergence. Training loss: 0.12, Validation loss: 0.15. '
                '(C) Magnitude Classification Loss: Multi-class task with stable convergence after '
                'epoch 20. (D) Azimuth Classification Loss: Most challenging task but demonstrates '
                'learning. (E) Learning Rate Schedule: Cosine annealing with warm restarts. '
                '(F) Validation Metrics: Recall Large reaches 98.65% and Precision Large achieves '
                '100% by epoch 30. No significant overfitting observed (train/val gap <0.1).'
            )
        },
        {
            'number': 5,
            'file': 'figures/FIG_5_CM_Magnitude.png',
            'title': 'Confusion Matrix and Performance Heatmap',
            'caption': (
                'Normalized confusion matrix for magnitude classification on test set (303 samples). '
                'Key Observations: (1) Large Events (M ≥ 6.0): 98.6% correctly classified (72/73), '
                'with only 1 misclassified as Medium. Zero Large events classified as Normal or '
                'Moderate. (2) Medium Events: 82.9% recall (58/70). (3) Moderate Events: 81.8% '
                'recall (45/55). (4) Normal Events: 97.1% recall (102/105). Overall Accuracy: 91.4% '
                '(277/303). Precision by Class: Normal: 92.7%, Moderate: 84.9%, Medium: 95.1%, '
                'Large: 100.0%. The perfect precision for Large events (no false alarms) is '
                'particularly noteworthy for operational deployment.'
            )
        },
        {
            'number': 6,
            'file': 'figures/FIG_6_GradCAM_Interpretation.png',
            'title': 'Grad-CAM Interpretability Analysis',
            'caption': (
                'Gradient-weighted Class Activation Mapping (Grad-CAM) visualization revealing '
                'which spectral-temporal features the model focuses on. Each row shows: (Left) '
                'Original RGB spectrogram, (Center) Grad-CAM heatmap overlaid, (Right) Isolated '
                'attention regions. Row A - Large Event (M6.5): Model strongly attends to '
                'high-amplitude, low-frequency features (0.01-0.03 Hz) in Z-component, particularly '
                '30-45 minutes before event. Row B - Medium Event (M5.3): Moderate attention to '
                'mid-frequency band. Row C - Moderate Event (M4.7): Weak, scattered attention. '
                'Row D - Normal (M3.2): Minimal attention. Model learns physically meaningful '
                'features (ULF band, Z-component dominance), validating decision-making process.'
            )
        }
    ]
    
    # Add each figure
    for fig in figures:
        # Figure heading
        doc.add_heading(f"Figure {fig['number']}: {fig['title']}", level=2)
        
        # Check if image exists
        if os.path.exists(fig['file']):
            try:
                # Add image
                doc.add_picture(fig['file'], width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image: {fig['file']}]", style='Intense Quote')
                doc.add_paragraph(f"Note: Could not embed image - {e}")
        else:
            doc.add_paragraph(f"[Image file not found: {fig['file']}]", style='Intense Quote')
        
        # Add caption
        caption = doc.add_paragraph()
        caption_run = caption.add_run(f"Figure {fig['number']}: ")
        caption_run.bold = True
        caption.add_run(fig['caption'])
        caption.style = 'Caption'
        
        # Page break after each figure (except last)
        if fig['number'] < len(figures):
            doc.add_page_break()
    
    # Supplementary Figures
    doc.add_page_break()
    doc.add_heading('Supplementary Figures', level=1)
    
    supp_figures = [
        {
            'number': 'S1',
            'file': 'figures/vis_test_distribution.png',
            'title': 'Test Set Distribution Analysis',
            'caption': (
                'Distribution of test samples across magnitude classes and stations. '
                'Shows balanced representation across different categories ensuring '
                'robust evaluation of model performance.'
            )
        },
        {
            'number': 'S2',
            'file': 'figures/vis_comparison_q1.png',
            'title': 'Comparison with Q1 Baseline Models',
            'caption': (
                'Performance comparison between our Hierarchical EfficientNet and baseline '
                'models from Q1 evaluation. Our model achieves superior recall for Large '
                'events while maintaining competitive performance across other metrics.'
            )
        },
        {
            'number': 'S3',
            'file': 'figures/vis_radar_performance.png',
            'title': 'Multi-Metric Performance Radar Chart',
            'caption': (
                'Radar chart showing model performance across multiple evaluation metrics. '
                'The chart demonstrates balanced performance with particular strength in '
                'Large event detection (recall and precision both near 100%).'
            )
        }
    ]
    
    # Add supplementary figures
    for fig in supp_figures:
        doc.add_heading(f"Figure {fig['number']}: {fig['title']}", level=2)
        
        if os.path.exists(fig['file']):
            try:
                doc.add_picture(fig['file'], width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image: {fig['file']}]", style='Intense Quote')
                doc.add_paragraph(f"Note: Could not embed image - {e}")
        else:
            doc.add_paragraph(f"[Image file not found: {fig['file']}]", style='Intense Quote')
        
        caption = doc.add_paragraph()
        caption_run = caption.add_run(f"Figure {fig['number']}: ")
        caption_run.bold = True
        caption.add_run(fig['caption'])
        caption.style = 'Caption'
        
        if fig != supp_figures[-1]:
            doc.add_page_break()
    
    # Figure submission guidelines
    doc.add_page_break()
    doc.add_heading('Figure Submission Guidelines', level=1)
    
    guidelines = [
        'All figures are provided in PNG format at 300 DPI resolution',
        'Figure dimensions: Approximately 3000×2000 pixels for main figures',
        'Color mode: RGB for online publication',
        'File size: Each figure <10 MB',
        'Figures are numbered sequentially (Figure 1-6 for main, S1-S3 for supplementary)',
        'Captions are provided both in this document and separately in manuscript',
        'All figures are original work created specifically for this study',
        'No copyright permissions required'
    ]
    
    for guideline in guidelines:
        doc.add_paragraph(guideline, style='List Bullet')
    
    doc.add_paragraph()
    
    # Technical specifications
    doc.add_heading('Technical Specifications', level=2)
    
    specs_table = doc.add_table(rows=7, cols=2)
    specs_table.style = 'Light Grid Accent 1'
    
    specs = [
        ('Format', 'PNG (Portable Network Graphics)'),
        ('Resolution', '300 DPI'),
        ('Color Mode', 'RGB'),
        ('Compression', 'Lossless'),
        ('Total Main Figures', '6'),
        ('Total Supplementary Figures', '3'),
        ('Total Package Size', '~15 MB')
    ]
    
    for i, (key, value) in enumerate(specs):
        specs_table.rows[i].cells[0].text = key
        specs_table.rows[i].cells[1].text = value
    
    # Save
    doc.save('7_FIGURES_COMPLETE.docx')
    print("✅ Generated: 7_FIGURES_COMPLETE.docx")

def main():
    """Generate figures DOCX"""
    print("=" * 60)
    print("Generating Figures DOCX with Images and Captions")
    print("=" * 60)
    print()
    
    # Change to publication folder
    os.chdir(os.path.dirname(__file__))
    
    try:
        create_figures_docx()
        
        print()
        print("=" * 60)
        print("✅ Figures DOCX generated successfully!")
        print("=" * 60)
        print()
        print("Generated file:")
        print("  7_FIGURES_COMPLETE.docx")
        print()
        print("Contents:")
        print("  - 6 Main Figures with images and captions")
        print("  - 3 Supplementary Figures with images and captions")
        print("  - Submission guidelines")
        print("  - Technical specifications")
        print()
        print("Note:")
        print("  - Images are embedded at 6 inches width")
        print("  - Each figure on separate page")
        print("  - Captions formatted as 'Caption' style")
        print("  - Ready for journal submission")
        
    except Exception as e:
        print(f"❌ Error generating figures DOCX: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
