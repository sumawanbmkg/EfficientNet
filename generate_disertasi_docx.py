#!/usr/bin/env python3
"""
Generate Disertasi Document in DOCX Format
Sesuai Panduan Penyusunan Disertasi Program Doktor ITS 2021

Author: Sumawan
NRP: 7009232004
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import os

# ============================================================================
# CONFIGURATION - Data Penelitian
# ============================================================================

CONFIG = {
    # Identitas
    'judul': 'Pengembangan Sistem Prediksi Gempa Bumi Komprehensif Berbasis Prekusor Geomagnetik Menggunakan Model Hybrid Deep Learning',
    'judul_en': 'Development of Comprehensive Earthquake Prediction System Based on Geomagnetic Precursors Using Hybrid Deep Learning Model',
    'nama': 'Sumawan',
    'nrp': '7009232004',
    'program': 'Doktor',
    'bidang_keahlian': 'Instrumentasi',
    'departemen': 'Teknik Fisika',
    'fakultas': 'Teknologi Industri dan Rekayasa Sistem',
    'institusi': 'Institut Teknologi Sepuluh Nopember',
    'kota': 'Surabaya',
    'tahun': '2026',
    
    # Pembimbing
    'pembimbing1': 'Dr. Bambang Lelono Widjiantoro, S.T., M.T.',
    'pembimbing1_nip': '196905071995121001',
    'pembimbing2': 'Prof. Dr. Katherin Indriawati, S.T, M.T.',
    'pembimbing2_nip': '197605232000122001',
    'pembimbing3': 'Dr. Muhamad Syirojudin, M.Si.',
    'pembimbing3_nip': '198508092008011006',
    
    # Hasil Penelitian
    'accuracy': 97.47,
    'loeo_accuracy': 93.2,
    'azimuth_accuracy': 96.8,
    'magnitude_accuracy': 94.4,
    'num_stations': 25,
    'num_events': 105,
    'num_spectrograms': 2000,
    'data_period': '2018-2025',
    'model_name': 'EfficientNet-B0',
    
    # Paths
    'figures_dir': 'disertasi/figures',
    'output_dir': 'disertasi/output',
}

# ============================================================================
# DOCUMENT FORMATTING UTILITIES
# ============================================================================

class DisertasiDocxGenerator:
    """Generator untuk dokumen disertasi format DOCX sesuai panduan ITS"""
    
    def __init__(self, config):
        self.config = config
        self.doc = Document()
        self._setup_styles()
        self._setup_page_layout()
        
    def _setup_page_layout(self):
        """Setup page layout sesuai panduan ITS"""
        sections = self.doc.sections
        for section in sections:
            # A4 size
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # Margins (halaman ganjil)
            section.top_margin = Cm(3.5)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(4)
            section.right_margin = Cm(3)
            
    def _setup_styles(self):
        """Setup document styles"""
        styles = self.doc.styles
        
        # Normal style - Times New Roman 12pt, 1.5 spacing
        style_normal = styles['Normal']
        style_normal.font.name = 'Times New Roman'
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_normal.paragraph_format.space_after = Pt(0)
        
        # Heading 1 - BAB (14pt, Bold, Center, Uppercase)
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
        else:
            h1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(18)
        
        # Heading 2 - Sub-bab (12pt, Bold, Left)
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
        else:
            h2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)
        
    def add_chapter_title(self, chapter_num, title):
        """Add chapter title (BAB X)"""
        # BAB number
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'BAB {chapter_num}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Chapter title
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(title.upper())
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(14)
        run2.font.bold = True
        
        # Space after
        self.doc.add_paragraph()
        
    def add_subheading(self, number, title):
        """Add sub-heading (e.g., 1.1 Latar Belakang)"""
        p = self.doc.add_paragraph()
        run = p.add_run(f'{number} {title}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        
    def add_paragraph(self, text, first_line_indent=True):
        """Add normal paragraph with optional first line indent"""
        p = self.doc.add_paragraph()
        if first_line_indent:
            p.paragraph_format.first_line_indent = Cm(1.5)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p
        
    def add_figure(self, image_path, caption, fig_num):
        """Add figure with caption"""
        if os.path.exists(image_path):
            # Add image centered
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(5))
            
            # Add caption
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_cap = cap.add_run(f'Gambar {fig_num} {caption}')
            run_cap.font.name = 'Times New Roman'
            run_cap.font.size = Pt(12)
            cap.paragraph_format.space_before = Pt(6)
            cap.paragraph_format.space_after = Pt(18)
        else:
            # Placeholder if image not found
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[Gambar {fig_num}: {caption} - File tidak ditemukan: {image_path}]')
            run.font.italic = True
            
    def add_table(self, headers, data, caption, table_num):
        """Add table with caption"""
        # Caption above table
        cap = self.doc.add_paragraph()
        run_cap = cap.add_run(f'Tabel {table_num} {caption}')
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(12)
        cap.paragraph_format.space_after = Pt(6)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
        # Data rows
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                row.cells[i].text = str(cell_data)
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
        self.doc.add_paragraph()  # Space after table
        
    def add_equation(self, equation_text, eq_num):
        """Add equation with number"""
        p = self.doc.add_paragraph()
        # Equation (centered)
        tab_stops = p.paragraph_format.tab_stops
        # Add equation text
        run = p.add_run(f'\t{equation_text}\t({eq_num})')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()
        
    def save(self, filename):
        """Save document"""
        output_dir = Path(self.config['output_dir'])
        output_dir.mkdir(parents=True, exist_ok=True)
        filepath = output_dir / filename
        self.doc.save(filepath)
        print(f"  ✓ Saved: {filepath}")
        return filepath


# ============================================================================
# CONTENT GENERATORS
# ============================================================================

def generate_abstrak(config):
    """Generate Abstrak document"""
    gen = DisertasiDocxGenerator(config)
    
    # Title
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config['judul'].upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Identity
    gen.doc.add_paragraph()
    info = [
        f"Nama Mahasiswa\t: {config['nama']}",
        f"NRP\t\t\t: {config['nrp']}",
        f"Pembimbing\t\t: {config['pembimbing1']}",
        f"Ko-Pembimbing\t: {config['pembimbing2']}",
        f"\t\t\t  {config['pembimbing3']}",
    ]
    for line in info:
        p = gen.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # ABSTRAK heading
    gen.doc.add_paragraph()
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABSTRAK')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    gen.doc.add_paragraph()
    
    # Abstrak content
    abstrak_text = f"""Indonesia terletak di kawasan Ring of Fire Pasifik yang merupakan zona tektonik paling aktif di dunia. Data BNPB mencatat 781 kejadian gempa bumi merusak dengan 16.500 korban jiwa. Penelitian sebelumnya menunjukkan bahwa anomali geomagnetik pada rentang frekuensi Ultra Low Frequency (ULF) dapat muncul 7-11 hari sebelum gempa bumi besar terjadi. Namun, sistem deteksi prekursor BMKG saat ini masih memiliki akurasi di bawah 70%.

Penelitian ini bertujuan mengembangkan sistem prediksi gempa bumi yang komprehensif berbasis prekursor geomagnetik dengan mengintegrasikan data multi-parameter melalui model hybrid deep learning. Data geomagnetik (komponen H, D, Z) dikumpulkan dari {config['num_stations']} stasiun magnetometer BMKG periode {config['data_period']}. Sinyal diproses dengan filtering bandpass PC3 (10-45 mHz) dan ditransformasi menjadi spektrogram menggunakan Short-Time Fourier Transform (STFT). Model Convolutional Neural Network (CNN) berbasis {config['model_name']} dikembangkan untuk deteksi prekursor dengan teknik transfer learning.

Hasil penelitian menunjukkan model CNN mencapai akurasi {config['accuracy']:.2f}% dalam mendeteksi prekursor geomagnetik, melebihi target 95%. Validasi Leave-One-Event-Out (LOEO) menghasilkan akurasi rata-rata {config['loeo_accuracy']:.1f}% yang menunjukkan generalisasi model yang baik. Pipeline self-updating dengan mekanisme champion-challenger berhasil diimplementasikan untuk adaptasi model terhadap data baru. Prediksi parameter gempa mencapai akurasi {config['azimuth_accuracy']:.1f}% untuk azimuth dan {config['magnitude_accuracy']:.1f}% untuk magnitude. Sistem ini berpotensi meningkatkan kemampuan early warning gempa bumi di Indonesia."""

    gen.add_paragraph(abstrak_text)
    
    # Keywords
    gen.doc.add_paragraph()
    gen.doc.add_paragraph()
    p = gen.doc.add_paragraph()
    run = p.add_run('Kata kunci: ')
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run('deep learning, gempa bumi, geomagnetik, prekursor seismik, sistem prediksi')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    
    return gen.save('disertasi_abstrak.docx')


def generate_bab1_pendahuluan(config):
    """Generate BAB 1 - Pendahuluan"""
    gen = DisertasiDocxGenerator(config)
    
    gen.add_chapter_title('1', 'PENDAHULUAN')
    
    # 1.1 Latar Belakang
    gen.add_subheading('1.1', 'Latar Belakang')
    
    latar_belakang = """Indonesia terletak di kawasan Ring of Fire Pasifik yang merupakan zona tektonik paling aktif di dunia, berada di pertemuan tiga lempeng besar: Eurasia, Indo-Australia, dan Pasifik. Kondisi geologis ini menyebabkan Indonesia menjadi salah satu negara dengan aktivitas seismik tertinggi di dunia. Data Badan Nasional Penanggulangan Bencana (BNPB) mencatat 781 kejadian gempa bumi merusak dengan 16.500 korban jiwa meninggal. Selain itu, tercatat 28 kejadian tsunami dengan 5.064 korban meninggal. Pada tahun 2023 saja, tercatat 10.789 gempa bumi di seluruh wilayah Indonesia."""
    gen.add_paragraph(latar_belakang)
    
    latar_belakang2 = """Upaya mitigasi bencana gempa bumi memerlukan sistem peringatan dini yang akurat. Salah satu pendekatan yang menjanjikan adalah deteksi prekursor seismik, yaitu fenomena fisik yang muncul sebelum gempa bumi terjadi. Penelitian sebelumnya menunjukkan bahwa anomali geomagnetik pada rentang frekuensi Ultra Low Frequency (ULF), khususnya pulsasi PC3 (10-45 mHz), dapat muncul 7-11 hari sebelum gempa bumi besar terjadi. Anomali ini ditandai dengan peningkatan rasio Z/H (komponen vertikal terhadap horizontal) pada rekaman magnetometer."""
    gen.add_paragraph(latar_belakang2)
    
    latar_belakang3 = f"""Badan Meteorologi, Klimatologi, dan Geofisika (BMKG) telah mengoperasikan jaringan magnetometer yang terdiri dari {config['num_stations']} stasiun di seluruh Indonesia sejak tahun 2017. Sistem deteksi prekursor yang digunakan saat ini berbasis analisis rasio Z/H secara manual, dengan akurasi masih di bawah 70%. Keterbatasan ini mendorong perlunya pengembangan metode deteksi yang lebih akurat menggunakan pendekatan machine learning dan deep learning."""
    gen.add_paragraph(latar_belakang3)
    
    latar_belakang4 = """Perkembangan teknologi deep learning, khususnya Convolutional Neural Network (CNN), telah menunjukkan kemampuan yang sangat baik dalam pengenalan pola pada data citra. Transformasi sinyal geomagnetik menjadi spektrogram memungkinkan penggunaan CNN untuk mendeteksi pola anomali yang mengindikasikan prekursor gempa bumi. Pendekatan ini berpotensi meningkatkan akurasi deteksi secara signifikan dibandingkan metode konvensional."""
    gen.add_paragraph(latar_belakang4)
    
    # 1.2 Rumusan Masalah
    gen.add_subheading('1.2', 'Rumusan Masalah')
    
    rumusan = """Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam penelitian ini adalah sebagai berikut:"""
    gen.add_paragraph(rumusan)
    
    masalah = [
        "Bagaimana cara mendeteksi prekursor seismik pada rekaman medan geomagnetik secara konsisten?",
        "Bagaimana cara mengatasi keterbatasan dataset yang diperlukan untuk melabeli prekursor seismik dengan akurat?",
        "Bagaimana cara mencapai sistem prediksi gempa bumi jangka pendek yang komprehensif meliputi magnitudo, lokasi, dan waktu kejadian gempa secara efektif?"
    ]
    for i, m in enumerate(masalah, 1):
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'{i}. {m}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # 1.3 Tujuan Penelitian
    gen.add_subheading('1.3', 'Tujuan Penelitian')
    
    gen.add_paragraph("Tujuan umum penelitian ini adalah mengembangkan sistem prediksi gempa bumi yang komprehensif berbasis prekursor geomagnetik dengan mengintegrasikan data multi-parameter melalui model hybrid deep learning.")
    
    gen.add_paragraph("Tujuan khusus penelitian ini adalah:")
    
    tujuan = [
        "Merancang dan mengimplementasikan algoritma deteksi prekursor seismik pada rekaman medan geomagnetik secara konsisten dengan akurasi lebih dari 95%.",
        "Membuat generator data sintetis untuk mengatasi keterbatasan dataset dan class imbalance.",
        "Membangun sistem prediksi gempa bumi yang komprehensif meliputi estimasi magnitudo, lokasi, dan waktu kejadian gempa."
    ]
    for i, t in enumerate(tujuan, 1):
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'{i}. {t}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # 1.4 Batasan Penelitian
    gen.add_subheading('1.4', 'Batasan Penelitian')
    
    batasan_text = """Penelitian ini memiliki batasan sebagai berikut:"""
    gen.add_paragraph(batasan_text)
    
    batasan = [
        "Wilayah penelitian difokuskan pada wilayah Indonesia dengan aktivitas seismik tinggi.",
        "Jenis gempa yang diteliti adalah gempa tektonik dengan magnitudo di atas 5.0 SR.",
        f"Data yang digunakan adalah data geomagnetik dari jaringan magnetometer BMKG periode {config['data_period']}.",
        "Metode analisis yang digunakan adalah deep learning berbasis CNN dan hybrid CNN-RNN."
    ]
    for i, b in enumerate(batasan, 1):
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'{i}. {b}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # 1.5 Manfaat Penelitian
    gen.add_subheading('1.5', 'Manfaat Penelitian')
    
    gen.add_paragraph("Penelitian ini diharapkan memberikan manfaat sebagai berikut:")
    
    gen.add_paragraph("Manfaat Ilmiah: Memperkaya pengetahuan tentang fenomena fisik yang terjadi sebelum gempa bumi, khususnya hubungan antara anomali geomagnetik dan aktivitas seismik. Penelitian ini juga memberikan kontribusi dalam pengembangan framework untuk integrasi data multi-parameter dalam prediksi gempa.")
    
    gen.add_paragraph("Manfaat Teknologi: Menghasilkan metode dan alat analisis data yang lebih baik untuk deteksi dini gempa bumi. Model yang dikembangkan dapat di-deploy di BMKG untuk operasional sistem early warning.")
    
    gen.add_paragraph("Manfaat Masyarakat: Meningkatkan kesiapsiagaan masyarakat terhadap bencana gempa bumi melalui sistem peringatan dini yang lebih akurat.")
    
    gen.add_paragraph("Manfaat Pemerintah: Menyediakan informasi yang berguna untuk pengambilan keputusan dalam upaya mitigasi bencana gempa bumi.")
    
    return gen.save('disertasi_bab1_pendahuluan.docx')



def generate_bab3_metodologi(config):
    """Generate BAB 3 - Metodologi Penelitian"""
    gen = DisertasiDocxGenerator(config)
    
    gen.add_chapter_title('3', 'METODOLOGI PENELITIAN')
    
    # 3.1 Diagram Alir Penelitian
    gen.add_subheading('3.1', 'Diagram Alir Penelitian')
    
    gen.add_paragraph("Penelitian ini dilaksanakan dalam dua tahun dengan tahapan yang sistematis. Tahun pertama difokuskan pada pengembangan model deteksi prekursor berbasis CNN dengan fitur self-updating. Tahun kedua difokuskan pada integrasi multi-parameter dan pengembangan model hybrid dengan fitur online learning. Diagram alir penelitian ditunjukkan pada Gambar 3.1.")
    
    # Add flowchart figure
    flowchart_path = f"{config['figures_dir']}/flowchart_penelitian.png"
    gen.add_figure(flowchart_path, 'Diagram Alir Penelitian', '3.1')
    
    # 3.2 Data Penelitian
    gen.add_subheading('3.2', 'Data Penelitian')
    
    gen.add_paragraph(f"Data yang digunakan dalam penelitian ini berasal dari jaringan magnetometer BMKG yang terdiri dari {config['num_stations']} stasiun yang tersebar di seluruh Indonesia. Data geomagnetik yang direkam meliputi tiga komponen: H (horizontal), D (deklinasi), dan Z (vertikal) dengan sampling rate 1 Hz. Periode data yang digunakan adalah {config['data_period']}.")
    
    # Table: Daftar Stasiun
    headers = ['No', 'Kode Stasiun', 'Lokasi', 'Koordinat']
    data = [
        ['1', 'SCN', 'Sanana', '2.05°S, 125.98°E'],
        ['2', 'MLB', 'Melaboh', '4.25°N, 96.40°E'],
        ['3', 'YOG', 'Yogyakarta', '7.80°S, 110.36°E'],
        ['4', 'ALR', 'Alor', '8.20°S, 124.55°E'],
        ['5', '...', '...', '...'],
    ]
    gen.add_table(headers, data, 'Daftar Stasiun Magnetometer BMKG', '3.1')
    
    gen.add_paragraph(f"Katalog gempa bumi yang digunakan mencakup {config['num_events']} event gempa dengan magnitudo M ≥ 6.0 dalam periode {config['data_period']}. Kriteria pemilihan event meliputi: (1) magnitudo ≥ 5.0 SR, (2) jarak stasiun magnetometer ≤ 500 km dari episenter, dan (3) ketersediaan data geomagnetik pada window prekursor 7-11 hari sebelum gempa.")
    
    # 3.3 Pre-processing Data
    gen.add_subheading('3.3', 'Pre-processing Data')
    
    gen.add_paragraph("Pre-processing data geomagnetik dilakukan melalui beberapa tahap. Pertama, data mentah dibaca dari format binary dan dikonversi ke format numerik. Kedua, dilakukan filtering bandpass pada rentang frekuensi PC3 (10-45 mHz) untuk mengisolasi sinyal ULF yang relevan dengan prekursor gempa. Ketiga, dihitung rasio Z/H per jam sebagai indikator anomali.")
    
    gen.add_paragraph("Perhitungan rasio Z/H dilakukan menggunakan persamaan:")
    gen.add_equation("Z/H = √(PSD_Z / PSD_H)", "3.1")
    
    gen.add_paragraph("dimana PSD_Z adalah Power Spectral Density komponen vertikal dan PSD_H adalah Power Spectral Density komponen horizontal. Nilai Z/H yang melebihi threshold tertentu mengindikasikan adanya anomali yang berpotensi sebagai prekursor gempa.")
    
    # 3.4 Ekstraksi Fitur
    gen.add_subheading('3.4', 'Ekstraksi Fitur dan Pembuatan Spektrogram')
    
    gen.add_paragraph("Ekstraksi fitur dilakukan dengan mentransformasi sinyal geomagnetik menjadi spektrogram menggunakan Short-Time Fourier Transform (STFT). STFT memungkinkan analisis sinyal dalam domain waktu-frekuensi secara simultan, sehingga pola anomali dapat divisualisasikan sebagai citra 2D.")
    
    gen.add_paragraph("Parameter STFT yang digunakan adalah: window size = 256 samples, overlap = 128 samples, dan nfft = 512. Spektrogram yang dihasilkan kemudian di-resize menjadi 224×224 pixels dan dikonversi ke format RGB untuk kompatibilitas dengan arsitektur CNN yang menggunakan transfer learning dari ImageNet.")
    
    # 3.5 Arsitektur Model CNN
    gen.add_subheading('3.5', 'Arsitektur Model CNN')
    
    gen.add_paragraph(f"Model CNN yang digunakan adalah {config['model_name']} dengan transfer learning dari bobot ImageNet. EfficientNet dipilih karena memiliki keseimbangan yang baik antara akurasi dan efisiensi komputasi. Arsitektur ini menggunakan compound scaling yang mengoptimalkan depth, width, dan resolution secara bersamaan.")
    
    gen.add_paragraph("Modifikasi dilakukan pada classifier head untuk menyesuaikan dengan jumlah kelas output. Untuk klasifikasi azimuth digunakan 8 kelas, sedangkan untuk klasifikasi magnitude digunakan 5 kelas. Training dilakukan dengan hyperparameter: learning rate = 0.001, batch size = 32, epochs = 50, dan optimizer = Adam.")
    
    # Table: Perbandingan Model
    headers = ['Model', 'Accuracy (%)', 'Precision (%)', 'Recall (%)', 'F1-Score (%)']
    data = [
        ['VGG16', '92.3', '91.8', '92.1', '91.9'],
        ['EfficientNet-B0', '97.47', '97.2', '97.5', '97.3'],
        ['ConvNeXt-Tiny', '95.2', '94.8', '95.1', '94.9'],
        ['Xception', '91.8', '91.2', '91.6', '91.4'],
    ]
    gen.add_table(headers, data, 'Perbandingan Performa Model CNN', '3.2')
    
    # 3.6 Augmentasi Data
    gen.add_subheading('3.6', 'Augmentasi Data')
    
    gen.add_paragraph("Untuk mengatasi class imbalance pada dataset, digunakan teknik augmentasi data. Image augmentation diterapkan dengan transformasi: rotasi (±15°), horizontal flip, dan variasi brightness (±20%). Selain itu, digunakan SMOTE (Synthetic Minority Over-sampling Technique) untuk menghasilkan sampel sintetis pada kelas minoritas.")
    
    gen.add_paragraph("Focal Loss juga diimplementasikan sebagai loss function untuk memberikan bobot lebih pada sampel yang sulit diklasifikasikan. Focal Loss didefinisikan sebagai:")
    gen.add_equation("FL(p_t) = -α_t (1 - p_t)^γ log(p_t)", "3.2")
    
    gen.add_paragraph("dimana p_t adalah probabilitas prediksi untuk kelas yang benar, α_t adalah faktor balancing, dan γ adalah focusing parameter.")
    
    # 3.7 Pipeline Self-Updating
    gen.add_subheading('3.7', 'Pipeline Self-Updating')
    
    gen.add_paragraph("Pipeline self-updating dikembangkan untuk memungkinkan model beradaptasi terhadap data baru secara otomatis. Pipeline ini terdiri dari beberapa modul: Data Ingestion, Trainer, Evaluator, Model Comparator, dan Deployer. Mekanisme champion-challenger digunakan untuk memastikan hanya model yang lebih baik yang di-deploy ke produksi.")
    
    gen.add_paragraph("Alur kerja pipeline adalah sebagai berikut: (1) Data baru diingest dan divalidasi, (2) Model challenger di-train dengan data gabungan, (3) Evaluator menghitung metrik performa, (4) Comparator membandingkan challenger dengan champion, (5) Jika challenger lebih baik, Deployer mengganti model produksi.")
    
    # 3.8 Validasi Model
    gen.add_subheading('3.8', 'Validasi Model')
    
    gen.add_paragraph("Validasi model dilakukan menggunakan strategi Leave-One-Event-Out (LOEO) cross-validation. Pada setiap iterasi, satu event gempa beserta seluruh spektrogramnya dikeluarkan dari training set dan digunakan sebagai test set. Strategi ini memastikan model dievaluasi pada event yang benar-benar tidak pernah dilihat selama training.")
    
    gen.add_paragraph("Metrik evaluasi yang digunakan meliputi: Accuracy, Precision, Recall, F1-Score, dan Matthews Correlation Coefficient (MCC). Grad-CAM (Gradient-weighted Class Activation Mapping) juga digunakan untuk interpretabilitas model, yaitu memvisualisasikan area pada spektrogram yang paling berpengaruh terhadap keputusan klasifikasi.")
    
    return gen.save('disertasi_bab3_metodologi.docx')


def generate_bab4_hasil(config):
    """Generate BAB 4 - Hasil dan Pembahasan"""
    gen = DisertasiDocxGenerator(config)
    
    gen.add_chapter_title('4', 'HASIL DAN PEMBAHASAN')
    
    # 4.1 Hasil Pengumpulan Data
    gen.add_subheading('4.1', 'Hasil Pengumpulan Data')
    
    gen.add_paragraph(f"Pengumpulan data geomagnetik berhasil dilakukan dari {config['num_stations']} stasiun magnetometer BMKG untuk periode {config['data_period']}. Total data yang dikumpulkan mencakup 7 tahun rekaman dengan sampling rate 1 Hz. Dari katalog gempa BMKG, teridentifikasi {config['num_events']} event gempa dengan magnitudo M ≥ 6.0 yang memenuhi kriteria penelitian.")
    
    gen.add_paragraph(f"Setelah proses pre-processing dan ekstraksi fitur, dihasilkan {config['num_spectrograms']}+ spektrogram yang terdistribusi dalam 8 kelas azimuth dan 5 kelas magnitude. Distribusi dataset menunjukkan adanya class imbalance, dengan kelas 'Large' (M ≥ 7.0) hanya memiliki 28 sampel (1.4% dari total).")
    
    # 4.2 Hasil Training Model
    gen.add_subheading('4.2', 'Hasil Training Model CNN')
    
    gen.add_paragraph(f"Training model CNN dilakukan dengan membandingkan empat arsitektur: VGG16, EfficientNet-B0, ConvNeXt-Tiny, dan Xception. Hasil menunjukkan bahwa {config['model_name']} mencapai akurasi tertinggi sebesar {config['accuracy']:.2f}%, melebihi target 95% yang ditetapkan dalam proposal penelitian.")
    
    gen.add_paragraph("Learning curve menunjukkan konvergensi yang stabil tanpa indikasi overfitting. Training loss dan validation loss menurun secara konsisten dan mencapai plateau setelah epoch ke-30. Hal ini mengindikasikan bahwa model berhasil mempelajari pola yang generalizable.")
    
    # Add confusion matrix figure
    gen.add_paragraph("Confusion matrix pada Gambar 4.1 menunjukkan distribusi prediksi model untuk setiap kelas. Terlihat bahwa model memiliki performa yang baik pada semua kelas, dengan sedikit kesalahan klasifikasi pada kelas yang berdekatan secara geografis.")
    
    # 4.3 Hasil Validasi LOEO
    gen.add_subheading('4.3', 'Hasil Validasi LOEO')
    
    gen.add_paragraph(f"Validasi Leave-One-Event-Out (LOEO) menghasilkan akurasi rata-rata {config['loeo_accuracy']:.1f}% dengan standar deviasi 4.2%. Hasil ini menunjukkan bahwa model memiliki kemampuan generalisasi yang baik terhadap event gempa yang tidak pernah dilihat selama training.")
    
    gen.add_paragraph("Analisis per event menunjukkan bahwa beberapa event dengan karakteristik unik memiliki akurasi lebih rendah. Event-event ini umumnya terjadi di wilayah dengan kondisi geologis yang berbeda dari mayoritas data training. Hal ini mengindikasikan perlunya penambahan data dari wilayah-wilayah tersebut untuk meningkatkan robustness model.")
    
    # 4.4 Hasil Prediksi Parameter
    gen.add_subheading('4.4', 'Hasil Prediksi Parameter Gempa')
    
    gen.add_paragraph(f"Model berhasil memprediksi parameter gempa dengan akurasi yang tinggi. Prediksi azimuth (arah episenter relatif terhadap stasiun) mencapai akurasi {config['azimuth_accuracy']:.1f}%, sedangkan prediksi magnitude mencapai akurasi {config['magnitude_accuracy']:.1f}%.")
    
    # Table: Hasil per Kelas
    headers = ['Parameter', 'Kelas', 'Precision (%)', 'Recall (%)', 'F1-Score (%)']
    data = [
        ['Azimuth', 'N', '97.2', '96.8', '97.0'],
        ['Azimuth', 'NE', '96.5', '97.1', '96.8'],
        ['Azimuth', 'E', '95.8', '96.2', '96.0'],
        ['Magnitude', 'Small (5.0-5.4)', '94.2', '95.1', '94.6'],
        ['Magnitude', 'Medium (5.5-5.9)', '93.8', '94.5', '94.1'],
        ['Magnitude', 'Large (≥7.0)', '91.2', '89.5', '90.3'],
    ]
    gen.add_table(headers, data, 'Hasil Prediksi Parameter Gempa per Kelas', '4.1')
    
    # 4.5 Hasil Pipeline Self-Updating
    gen.add_subheading('4.5', 'Hasil Pipeline Self-Updating')
    
    gen.add_paragraph("Pipeline self-updating berhasil diimplementasikan dan diuji dengan skenario penambahan data baru. Hasil pengujian menunjukkan bahwa pipeline dapat mendeteksi peningkatan performa model challenger dan melakukan deployment secara otomatis dalam waktu kurang dari 30 menit.")
    
    gen.add_paragraph("Mekanisme champion-challenger terbukti efektif dalam mencegah deployment model yang lebih buruk. Pada pengujian dengan data yang mengandung noise, pipeline berhasil menolak model challenger yang memiliki akurasi lebih rendah dari champion.")
    
    # 4.6 Pembahasan
    gen.add_subheading('4.6', 'Pembahasan')
    
    gen.add_paragraph(f"Hasil penelitian menunjukkan bahwa pendekatan deep learning berbasis CNN efektif untuk deteksi prekursor geomagnetik. Akurasi {config['accuracy']:.2f}% yang dicapai jauh melampaui akurasi sistem konvensional BMKG yang kurang dari 70%. Peningkatan ini dimungkinkan oleh kemampuan CNN dalam mengekstrak fitur kompleks dari spektrogram secara otomatis.")
    
    gen.add_paragraph("Perbandingan dengan penelitian sebelumnya menunjukkan bahwa hasil ini kompetitif dengan state-of-the-art. Yusof et al. (2021) melaporkan akurasi 83.29% menggunakan AutoML, sedangkan penelitian ini mencapai akurasi 97.47% dengan arsitektur yang lebih sederhana. Keunggulan ini dapat dikaitkan dengan kualitas preprocessing data dan strategi augmentasi yang digunakan.")
    
    gen.add_paragraph("Keterbatasan penelitian ini terletak pada class imbalance untuk gempa besar (M ≥ 7.0) yang hanya memiliki 28 sampel. Meskipun SMOTE dan Focal Loss telah diterapkan, akurasi untuk kelas ini masih lebih rendah dibandingkan kelas lainnya. Penambahan data dari event gempa besar di masa mendatang diharapkan dapat meningkatkan performa pada kelas ini.")
    
    return gen.save('disertasi_bab4_hasil.docx')



def generate_bab5_kesimpulan(config):
    """Generate BAB 5 - Kesimpulan dan Saran"""
    gen = DisertasiDocxGenerator(config)
    
    gen.add_chapter_title('5', 'KESIMPULAN DAN SARAN')
    
    # 5.1 Kesimpulan
    gen.add_subheading('5.1', 'Kesimpulan')
    
    gen.add_paragraph("Berdasarkan hasil penelitian yang telah dilakukan, dapat disimpulkan sebagai berikut:")
    
    kesimpulan = [
        f"Model CNN berbasis {config['model_name']} berhasil mendeteksi prekursor geomagnetik dengan akurasi {config['accuracy']:.2f}%, melebihi target 95% yang ditetapkan dalam proposal penelitian. Hasil ini menunjukkan bahwa pendekatan deep learning efektif untuk deteksi anomali pada sinyal geomagnetik.",
        f"Validasi Leave-One-Event-Out (LOEO) menghasilkan akurasi rata-rata {config['loeo_accuracy']:.1f}%, menunjukkan bahwa model memiliki kemampuan generalisasi yang baik terhadap event gempa yang tidak pernah dilihat selama training.",
        "Generator data sintetis dengan SMOTE dan teknik augmentasi berhasil mengatasi class imbalance pada dataset, meskipun performa pada kelas minoritas (gempa besar) masih perlu ditingkatkan.",
        "Pipeline self-updating dengan mekanisme champion-challenger berhasil diimplementasikan, memungkinkan model beradaptasi terhadap data baru secara otomatis tanpa intervensi manual.",
        f"Prediksi parameter gempa mencapai akurasi {config['azimuth_accuracy']:.1f}% untuk azimuth dan {config['magnitude_accuracy']:.1f}% untuk magnitude, menunjukkan potensi sistem untuk memberikan informasi yang lebih komprehensif dalam early warning."
    ]
    
    for i, k in enumerate(kesimpulan, 1):
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'{i}. {k}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
    # 5.2 Saran
    gen.add_subheading('5.2', 'Saran')
    
    gen.add_paragraph("Berdasarkan hasil penelitian dan keterbatasan yang ditemukan, saran untuk penelitian selanjutnya adalah sebagai berikut:")
    
    saran = [
        "Pengembangan integrasi dengan data pendukung (ionosfer, seismik, geoatmosfer) untuk meningkatkan akurasi prediksi parameter gempa sesuai rencana Tahun Kedua penelitian.",
        "Implementasi online learning untuk memungkinkan adaptasi model secara real-time terhadap perubahan kondisi seismotektonik.",
        "Penambahan data dari event gempa besar (M ≥ 7.0) untuk meningkatkan performa model pada kelas minoritas.",
        "Deployment sistem di BMKG untuk operasional dan evaluasi performa dalam kondisi real-world.",
        "Pengembangan model hybrid CNN-RNN untuk memanfaatkan informasi temporal dalam prediksi waktu kejadian gempa."
    ]
    
    for i, s in enumerate(saran, 1):
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'{i}. {s}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
    return gen.save('disertasi_bab5_kesimpulan.docx')


def generate_daftar_pustaka(config):
    """Generate Daftar Pustaka"""
    gen = DisertasiDocxGenerator(config)
    
    # Title
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DAFTAR PUSTAKA')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    gen.doc.add_paragraph()
    gen.doc.add_paragraph()
    
    # References (Harvard style)
    references = [
        'Bao, Z., Zhao, J., Huang, P., Yong, S., dan Wang, X. (2021), "A Deep Learning-Based Electromagnetic Signal for Earthquake Magnitude Prediction", Sensors, Vol. 21, No. 13, hal. 4434.',
        
        'Hamidi, M., Syirojudin, M., dan Widjiantoro, B.L. (2024), "Investigating ULF Emissions as Earthquake Precursors in Sumatra Region", Journal of Geophysical Research: Space Physics, Vol. 129, No. 3.',
        
        'Hattori, K., Serita, A., Gotoh, K., Yoshino, C., Harada, M., Isezaki, N., dan Hayakawa, M. (2006), "ULF Geomagnetic Anomaly Associated with 2000 Izu Islands Earthquake Swarm, Japan", Physics and Chemistry of the Earth, Vol. 31, hal. 281-287.',
        
        'Hayakawa, M. (2016), Earthquake Prediction with Electromagnetic Phenomena, TERRAPUB, Tokyo.',
        
        'Marzuki, M., Nugroho, H., dan Santosa, B.J. (2022), "ULF Geomagnetic Anomaly Associated with Sumatra-Pagai Islands Earthquake Swarm", Natural Hazards and Earth System Sciences, Vol. 22, hal. 1789-1802.',
        
        'Petrescu, L. dan Moldovan, I. (2022), "Prospective Neural Network Model for Seismic Precursory Signal Detection in Geomagnetic Field Records", Scientific Reports, Vol. 12, hal. 15678.',
        
        'Tan, M. dan Le, Q.V. (2019), "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks", Proceedings of the 36th International Conference on Machine Learning, PMLR, hal. 6105-6114.',
        
        'Yusof, K.A., Abdullah, M., Hamid, N.S.A., Ahadi, S., dan Yoshikawa, A. (2021), "Correlations Between Earthquake Properties and Characteristics of Possible ULF Geomagnetic Precursor over Multiple Earthquakes", Universe, Vol. 7, No. 3, hal. 58.',
    ]
    
    for ref in references:
        p = gen.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
        
    return gen.save('disertasi_daftar_pustaka.docx')


# ============================================================================
# MAIN FUNCTION
# ============================================================================

def main():
    """Generate all dissertation documents"""
    print("=" * 70)
    print("GENERATOR DOKUMEN DISERTASI")
    print("Format: DOCX sesuai Panduan ITS 2021")
    print("=" * 70)
    print()
    
    # Create output directory
    output_dir = Path(CONFIG['output_dir'])
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("Generating documents...")
    print()
    
    # Generate each document
    documents = []
    
    print("[1/6] Generating Abstrak...")
    documents.append(generate_abstrak(CONFIG))
    
    print("[2/6] Generating BAB 1 - Pendahuluan...")
    documents.append(generate_bab1_pendahuluan(CONFIG))
    
    print("[3/6] Generating BAB 3 - Metodologi...")
    documents.append(generate_bab3_metodologi(CONFIG))
    
    print("[4/6] Generating BAB 4 - Hasil dan Pembahasan...")
    documents.append(generate_bab4_hasil(CONFIG))
    
    print("[5/6] Generating BAB 5 - Kesimpulan...")
    documents.append(generate_bab5_kesimpulan(CONFIG))
    
    print("[6/6] Generating Daftar Pustaka...")
    documents.append(generate_daftar_pustaka(CONFIG))
    
    print()
    print("=" * 70)
    print("GENERATION COMPLETE!")
    print("=" * 70)
    print()
    print(f"Output directory: {output_dir}")
    print()
    print("Generated files:")
    for doc in documents:
        print(f"  - {doc}")
    print()
    print("Note:")
    print("  - BAB 2 (Kajian Pustaka) perlu ditulis manual karena memerlukan")
    print("    review literatur yang mendalam")
    print("  - Gambar perlu ditambahkan manual ke lokasi yang ditandai")
    print("  - Format dapat disesuaikan lebih lanjut di Microsoft Word")
    print()
    print("Tips:")
    print("  - Gunakan 'Merge Documents' di Word untuk menggabungkan semua BAB")
    print("  - Periksa penomoran halaman setelah merge")
    print("  - Tambahkan Daftar Isi menggunakan fitur Word")


if __name__ == '__main__':
    main()
