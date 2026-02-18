#!/usr/bin/env python3
"""
Generate comprehensive Supplementary Materials in Word format.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.add_run('SUPPLEMENTARY MATERIALS').bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)

subtitle = doc.add_paragraph()
subtitle.add_run('Earthquake Precursor Detection using Deep Learning: A Comparative Study of VGG16 and EfficientNet-B0 for Geomagnetic Anomaly Classification')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)

def add_heading(text):
    h = doc.add_paragraph()
    h.add_run(text).bold = True
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

# ============================================================================
# S1. GEOMAGNETIC STATION NETWORK
# ============================================================================
add_heading('S1. Geomagnetic Station Network')

doc.add_paragraph('Table S1 presents the coordinates of 24 geomagnetic stations operated by BMKG across the Indonesian archipelago used in this study.')

# Station table
table_s1 = doc.add_table(rows=25, cols=4)
table_s1.style = 'Table Grid'

headers = ['Station Code', 'Latitude (°)', 'Longitude (°)', 'Region']
for i, h in enumerate(headers):
    table_s1.rows[0].cells[i].text = h
    table_s1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

stations = [
    ('SBG', '5.877', '95.338', 'Aceh'),
    ('SCN', '-0.546', '100.298', 'West Sumatra'),
    ('KPY', '-3.680', '102.582', 'Bengkulu'),
    ('LWA', '-5.017', '104.058', 'Lampung'),
    ('LPS', '-5.789', '105.583', 'Banten'),
    ('SRG', '-6.171', '106.051', 'West Java'),
    ('SKB', '-7.074', '106.531', 'West Java'),
    ('CLP', '-7.719', '109.015', 'Central Java'),
    ('YOG', '-7.731', '110.354', 'Yogyakarta'),
    ('TRT', '-7.705', '112.635', 'East Java'),
    ('LUT', '-8.220', '116.407', 'West Nusa Tenggara'),
    ('ALR', '-8.144', '124.590', 'East Nusa Tenggara'),
    ('SMI', '-7.669', '131.579', 'Maluku'),
    ('SRO', '-0.863', '131.259', 'West Papua'),
    ('TNT', '0.813', '127.367', 'North Maluku'),
    ('TND', '1.295', '124.925', 'North Sulawesi'),
    ('GTO', '0.556', '123.141', 'Gorontalo'),
    ('LWK', '-1.000', '122.784', 'Central Sulawesi'),
    ('PLU', '-0.620', '119.859', 'Central Sulawesi'),
    ('TRD', '2.136', '117.424', 'East Kalimantan'),
    ('JYP', '-2.514', '140.704', 'Papua'),
    ('AMB', '-3.676', '128.111', 'Maluku'),
    ('GSI', '1.304', '97.576', 'North Sumatra'),
    ('MLB', '4.049', '96.248', 'Aceh'),
]

for i, (code, lat, lon, region) in enumerate(stations):
    row = table_s1.rows[i + 1]
    row.cells[0].text = code
    row.cells[1].text = lat
    row.cells[2].text = lon
    row.cells[3].text = region

doc.add_paragraph()

# ============================================================================
# S2. DATASET STATISTICS
# ============================================================================
add_heading('S2. Dataset Statistics')

doc.add_paragraph('Table S2 summarizes the dataset composition and class distribution.')

table_s2 = doc.add_table(rows=12, cols=3)
table_s2.style = 'Table Grid'

headers2 = ['Category', 'Value', 'Percentage']
for i, h in enumerate(headers2):
    table_s2.rows[0].cells[i].text = h
    table_s2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

dataset_stats = [
    ('Total Samples', '1,972', '100%'),
    ('Precursor Samples', '1,084', '55.0%'),
    ('Normal Samples', '888', '45.0%'),
    ('Unique Earthquake Events', '256', '-'),
    ('Training Set', '1,336', '67.7%'),
    ('Validation Set', '352', '17.9%'),
    ('Test Set', '284', '14.4%'),
    ('Temporal Windowing Factor', '4.2×', '-'),
    ('Image Size', '224×224 pixels', '-'),
    ('Color Channels', '3 (RGB)', '-'),
    ('Data Period', '2018-2025', '-'),
]

for i, (cat, val, pct) in enumerate(dataset_stats):
    row = table_s2.rows[i + 1]
    row.cells[0].text = cat
    row.cells[1].text = val
    row.cells[2].text = pct

doc.add_paragraph()

# Magnitude distribution
add_heading('S2.1 Magnitude Class Distribution')

table_mag = doc.add_table(rows=5, cols=4)
table_mag.style = 'Table Grid'

headers_mag = ['Class', 'Magnitude Range', 'Events', 'Percentage']
for i, h in enumerate(headers_mag):
    table_mag.rows[0].cells[i].text = h
    table_mag.rows[0].cells[i].paragraphs[0].runs[0].bold = True

mag_data = [
    ('Small', 'M4.0-4.9', '89', '34.8%'),
    ('Medium', 'M5.0-5.9', '112', '43.8%'),
    ('Large', 'M6.0-6.9', '42', '16.4%'),
    ('Major', 'M7.0+', '13', '5.1%'),
]

for i, row_data in enumerate(mag_data):
    for j, cell_data in enumerate(row_data):
        table_mag.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

# Azimuth distribution
add_heading('S2.2 Azimuth Class Distribution')

table_azi = doc.add_table(rows=10, cols=4)
table_azi.style = 'Table Grid'

headers_azi = ['Class', 'Direction Range', 'Samples', 'Percentage']
for i, h in enumerate(headers_azi):
    table_azi.rows[0].cells[i].text = h
    table_azi.rows[0].cells[i].paragraphs[0].runs[0].bold = True

azi_data = [
    ('N', '337.5°-22.5°', '28', '2.6%'),
    ('NE', '22.5°-67.5°', '35', '3.2%'),
    ('E', '67.5°-112.5°', '31', '2.9%'),
    ('SE', '112.5°-157.5°', '29', '2.7%'),
    ('S', '157.5°-202.5°', '33', '3.1%'),
    ('SW', '202.5°-247.5°', '27', '2.5%'),
    ('W', '247.5°-292.5°', '30', '2.8%'),
    ('NW', '292.5°-337.5°', '43', '4.0%'),
    ('Normal', 'No precursor', '888', '45.0%'),
]

for i, row_data in enumerate(azi_data):
    for j, cell_data in enumerate(row_data):
        table_azi.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

# ============================================================================
# S3. HYPERPARAMETER CONFIGURATION
# ============================================================================
add_heading('S3. Hyperparameter Configuration')

doc.add_paragraph('Table S3 presents the complete hyperparameter configuration for both models.')

table_hp = doc.add_table(rows=19, cols=3)
table_hp.style = 'Table Grid'

headers_hp = ['Hyperparameter', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_hp):
    table_hp.rows[0].cells[i].text = h
    table_hp.rows[0].cells[i].paragraphs[0].runs[0].bold = True

hp_data = [
    ('Pre-trained Weights', 'ImageNet', 'ImageNet'),
    ('Input Size', '224×224×3', '224×224×3'),
    ('Optimizer', 'Adam', 'Adam'),
    ('Learning Rate', '1×10⁻⁴', '9.89×10⁻⁴'),
    ('Learning Rate (Fine-tune)', '1×10⁻⁵', '1×10⁻⁵'),
    ('Weight Decay', '0', '5.2×10⁻⁵'),
    ('Batch Size', '32', '32'),
    ('Max Epochs', '30', '30'),
    ('Actual Epochs', '11', '12'),
    ('Best Epoch', '11', '5'),
    ('Dropout Rate', '0.5', '0.444'),
    ('Dense Units', '4096→512', '1280→512'),
    ('Early Stopping Patience', '5', '5'),
    ('Loss Function', 'CrossEntropy', 'CrossEntropy'),
    ('Class Weighting', 'None', 'None'),
    ('Data Augmentation', 'None', 'None'),
    ('Normalization', 'ImageNet stats', 'ImageNet stats'),
    ('Tuning Method', 'Manual', 'Optuna (20 trials)'),
]

for i, row_data in enumerate(hp_data):
    for j, cell_data in enumerate(row_data):
        table_hp.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

# ============================================================================
# S4. DETAILED RESULTS
# ============================================================================
add_heading('S4. Detailed Training and Test Results')

add_heading('S4.1 Training History')

table_train = doc.add_table(rows=9, cols=3)
table_train.style = 'Table Grid'

headers_train = ['Metric', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_train):
    table_train.rows[0].cells[i].text = h
    table_train.rows[0].cells[i].paragraphs[0].runs[0].bold = True

train_data = [
    ('Training Duration', '2.3 hours', '3.8 hours'),
    ('Final Training Loss', '0.2918', '1.0549'),
    ('Final Validation Loss', '2.8947', '0.8523'),
    ('Train-Val Loss Gap', '2.6029', '-0.2026'),
    ('Train Magnitude Acc', '99.64%', '97.01%'),
    ('Val Magnitude Acc', '97.18%', '97.73%'),
    ('Train Azimuth Acc', '93.14%', '69.09%'),
    ('Val Azimuth Acc', '59.51%', '77.27%'),
]

for i, row_data in enumerate(train_data):
    for j, cell_data in enumerate(row_data):
        table_train.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

add_heading('S4.2 Test Set Performance')

table_test = doc.add_table(rows=7, cols=3)
table_test.style = 'Table Grid'

headers_test = ['Metric', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_test):
    table_test.rows[0].cells[i].text = h
    table_test.rows[0].cells[i].paragraphs[0].runs[0].bold = True

test_data = [
    ('Test Samples', '284', '284'),
    ('Magnitude Accuracy', '98.68%', '94.37%'),
    ('Azimuth Accuracy', '54.93%', '57.39%'),
    ('Combined Accuracy', '76.81%', '75.88%'),
    ('Normal Detection', '100%', '100%'),
    ('Mag Train-Test Gap', '0.96%', '2.64%'),
]

for i, row_data in enumerate(test_data):
    for j, cell_data in enumerate(row_data):
        table_test.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

add_heading('S4.3 Per-Class Magnitude Metrics')

table_class = doc.add_table(rows=5, cols=7)
table_class.style = 'Table Grid'

headers_class = ['Class', 'VGG16 P', 'VGG16 R', 'VGG16 F1', 'EffNet P', 'EffNet R', 'EffNet F1']
for i, h in enumerate(headers_class):
    table_class.rows[0].cells[i].text = h
    table_class.rows[0].cells[i].paragraphs[0].runs[0].bold = True

class_data = [
    ('Small', '0.98', '0.96', '0.97', '0.94', '0.92', '0.93'),
    ('Medium', '0.96', '0.96', '0.96', '0.92', '0.92', '0.92'),
    ('Large', '0.91', '0.95', '0.93', '0.84', '0.88', '0.86'),
    ('Major', '0.92', '0.92', '0.92', '0.85', '0.85', '0.85'),
]

for i, row_data in enumerate(class_data):
    for j, cell_data in enumerate(row_data):
        table_class.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

# ============================================================================
# S5. MODEL ARCHITECTURE DETAILS
# ============================================================================
add_heading('S5. Model Architecture Details')

add_heading('S5.1 Architecture Comparison')

table_arch = doc.add_table(rows=13, cols=3)
table_arch.style = 'Table Grid'

headers_arch = ['Component', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_arch):
    table_arch.rows[0].cells[i].text = h
    table_arch.rows[0].cells[i].paragraphs[0].runs[0].bold = True

arch_data = [
    ('Total Parameters', '138 million', '5.3 million'),
    ('Model Size', '528 MB', '20 MB'),
    ('Convolutional Layers', '13', '8 MBConv blocks'),
    ('First Conv Filters', '64', '32'),
    ('Last Conv Filters', '512', '1280'),
    ('Activation Function', 'ReLU', 'SiLU (Swish)'),
    ('Squeeze-Excitation', 'No', 'Yes'),
    ('Depthwise Separable', 'No', 'Yes'),
    ('Skip Connections', 'No', 'Yes'),
    ('Pooling', 'MaxPool2d', 'Adaptive'),
    ('Estimated FLOPs', '15.5 GFLOPs', '0.39 GFLOPs'),
    ('Inference Time', '125 ms', '50 ms'),
]

for i, row_data in enumerate(arch_data):
    for j, cell_data in enumerate(row_data):
        table_arch.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

add_heading('S5.2 Multi-Task Head Architecture')

doc.add_paragraph('Both models share the same multi-task head architecture:')

head_desc = doc.add_paragraph()
head_desc.add_run('• Shared Dense Layer: ').bold = True
head_desc.add_run('512 units with ReLU activation and 50% dropout\n')
head_desc.add_run('• Magnitude Head: ').bold = True
head_desc.add_run('Linear layer (512 → 4) with softmax activation\n')
head_desc.add_run('• Azimuth Head: ').bold = True
head_desc.add_run('Linear layer (512 → 9) with softmax activation\n')
head_desc.add_run('• Loss Combination: ').bold = True
head_desc.add_run('Sum of magnitude and azimuth cross-entropy losses')

doc.add_paragraph()

# ============================================================================
# S6. GRAD-CAM ANALYSIS
# ============================================================================
add_heading('S6. Grad-CAM Explainability Analysis')

doc.add_paragraph('Table S6 summarizes the Grad-CAM analysis results for both models.')

table_gradcam = doc.add_table(rows=8, cols=3)
table_gradcam.style = 'Table Grid'

headers_gc = ['Aspect', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_gc):
    table_gradcam.rows[0].cells[i].text = h
    table_gradcam.rows[0].cells[i].paragraphs[0].runs[0].bold = True

gc_data = [
    ('Target Layer', 'Last MaxPool', 'Last Conv2dNormActivation'),
    ('Attention Pattern', 'Concentrated', 'Distributed'),
    ('ULF Band Focus', 'Yes (0.001-0.01 Hz)', 'Yes (0.001-0.01 Hz)'),
    ('Temporal Focus', '12-24h before event', '12-24h before event'),
    ('Prediction Agreement', '100% (3/3 samples)', '100% (3/3 samples)'),
    ('Average Confidence', '64.98%', '91.84%'),
    ('Physical Interpretability', 'High', 'High'),
]

for i, row_data in enumerate(gc_data):
    for j, cell_data in enumerate(row_data):
        table_gradcam.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

doc.add_paragraph('Key observations from Grad-CAM analysis:')
obs = doc.add_paragraph()
obs.add_run('1. ').bold = True
obs.add_run('Both models consistently focus on the ULF frequency band (0.001-0.01 Hz), validating physical interpretability.\n')
obs.add_run('2. ').bold = True
obs.add_run('Temporal attention patterns show increased activation 12-24 hours before earthquakes.\n')
obs.add_run('3. ').bold = True
obs.add_run('100% prediction agreement between models on analyzed samples suggests learned features are robust.\n')
obs.add_run('4. ').bold = True
obs.add_run('EfficientNet-B0 shows higher average confidence (91.84% vs 64.98%), indicating more decisive predictions.')

doc.add_paragraph()

# ============================================================================
# S7. COMPUTATIONAL REQUIREMENTS
# ============================================================================
add_heading('S7. Computational Requirements')

table_comp = doc.add_table(rows=8, cols=3)
table_comp.style = 'Table Grid'

headers_comp = ['Resource', 'VGG16', 'EfficientNet-B0']
for i, h in enumerate(headers_comp):
    table_comp.rows[0].cells[i].text = h
    table_comp.rows[0].cells[i].paragraphs[0].runs[0].bold = True

comp_data = [
    ('GPU Memory (Training)', '~8 GB', '~3 GB'),
    ('GPU Memory (Inference)', '~2 GB', '~500 MB'),
    ('CPU RAM', '16 GB', '8 GB'),
    ('Storage (Model)', '528 MB', '20 MB'),
    ('Training Time', '2.3 hours', '3.8 hours'),
    ('Inference Speed', '125 ms/sample', '50 ms/sample'),
    ('Throughput', '~8 samples/sec', '~20 samples/sec'),
]

for i, row_data in enumerate(comp_data):
    for j, cell_data in enumerate(row_data):
        table_comp.rows[i + 1].cells[j].text = cell_data

doc.add_paragraph()

# ============================================================================
# S8. CODE AVAILABILITY
# ============================================================================
add_heading('S8. Code and Data Availability')

doc.add_paragraph('All code, pre-trained models, and documentation are publicly available:')

avail = doc.add_paragraph()
avail.add_run('Repository: ').bold = True
avail.add_run('https://github.com/sumawanbmkg/earthquake-precursor-cnn\n\n')
avail.add_run('Contents:\n')
avail.add_run('• src/ - Source code for training, evaluation, and inference\n')
avail.add_run('• models/ - Pre-trained model download instructions\n')
avail.add_run('• notebooks/ - Jupyter notebooks for data exploration and visualization\n')
avail.add_run('• figures/ - All paper figures in high resolution\n')
avail.add_run('• docs/ - API documentation and guides\n\n')
avail.add_run('To reproduce results:\n')

code_block = doc.add_paragraph()
code_block.add_run('git clone https://github.com/sumawanbmkg/earthquake-precursor-cnn.git\n')
code_block.add_run('cd earthquake-precursor-cnn\n')
code_block.add_run('pip install -r requirements.txt\n')
code_block.add_run('python scripts/download_models.py\n')
code_block.add_run('python src/evaluate.py --model efficientnet')
code_block.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()
doc.add_paragraph('— End of Supplementary Materials —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = Path('publication_package/SUPPLEMENTARY_MATERIALS.docx')
doc.save(output_path)

print("=" * 60)
print("Supplementary Materials (Word) Generated Successfully!")
print("=" * 60)
print(f"Output: {output_path.absolute()}")
print()
print("Contents:")
print("  S1. Geomagnetic Station Network (24 stations)")
print("  S2. Dataset Statistics")
print("      S2.1 Magnitude Class Distribution")
print("      S2.2 Azimuth Class Distribution")
print("  S3. Hyperparameter Configuration")
print("  S4. Detailed Training and Test Results")
print("      S4.1 Training History")
print("      S4.2 Test Set Performance")
print("      S4.3 Per-Class Magnitude Metrics")
print("  S5. Model Architecture Details")
print("      S5.1 Architecture Comparison")
print("      S5.2 Multi-Task Head Architecture")
print("  S6. Grad-CAM Explainability Analysis")
print("  S7. Computational Requirements")
print("  S8. Code and Data Availability")
print("=" * 60)
