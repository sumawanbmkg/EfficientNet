#!/usr/bin/env python3
"""
Generate Supplementary Materials Word Document for ConvNeXt Paper
"""

import json
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("publication_convnext")
FIGURES_DIR = OUTPUT_DIR / "figures"
LOEO_RESULTS = Path("loeo_convnext_results/loeo_convnext_final_results.json")

def load_results():
    if LOEO_RESULTS.exists():
        with open(LOEO_RESULTS, 'r') as f:
            return json.load(f)
    return None

def create_supplementary():
    results = load_results()
    doc = Document()
    
    # Title
    title = doc.add_heading('Supplementary Materials', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('ConvNeXt for Earthquake Precursor Detection')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # S1: Dataset Details
    doc.add_heading('S1. Dataset Details', level=1)
    doc.add_paragraph("""Total spectrograms: 1,972
Earthquake events: 256
Time period: 2018-2025
Stations: SCN, MLB, GTO, TRD, TRT, LUT, SBG, SKB, and others""")

    # S2: Training Configuration
    doc.add_heading('S2. Training Configuration', level=1)
    config_text = """Model: ConvNeXt-Tiny
Parameters: 28,615,789
Pretrained: ImageNet-1K
Optimizer: AdamW
Learning Rate: 0.0001
Weight Decay: 0.05
Batch Size: 16
Epochs: 15 (with early stopping)
Scheduler: Cosine Annealing
Dropout: 0.5"""
    doc.add_paragraph(config_text)
    
    # S3: LOEO Results
    doc.add_heading('S3. LOEO Cross-Validation Results', level=1)
    
    if results:
        # Per-fold table
        table = doc.add_table(rows=12, cols=5)
        table.style = 'Table Grid'
        
        headers = ['Fold', 'Magnitude Acc', 'Azimuth Acc', 'Combined Acc', 'Test Samples']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        for idx, r in enumerate(results['per_fold_results']):
            row = table.rows[idx + 1]
            row.cells[0].text = str(r['fold'])
            row.cells[1].text = f"{r['magnitude_accuracy']:.2f}%"
            row.cells[2].text = f"{r['azimuth_accuracy']:.2f}%"
            row.cells[3].text = f"{r['combined_accuracy']:.2f}%"
            row.cells[4].text = str(r['n_test_samples'])
        
        # Summary row
        combined = [r['combined_accuracy'] for r in results['per_fold_results']]
        row = table.rows[11]
        row.cells[0].text = 'Mean ± Std'
        row.cells[1].text = f"{results['magnitude_accuracy']['mean']:.2f}% ± {results['magnitude_accuracy']['std']:.2f}%"
        row.cells[2].text = f"{results['azimuth_accuracy']['mean']:.2f}% ± {results['azimuth_accuracy']['std']:.2f}%"
        row.cells[3].text = f"{np.mean(combined):.2f}% ± {np.std(combined):.2f}%"
        row.cells[4].text = '-'
        
        doc.add_paragraph()
        doc.add_paragraph('Table S1: LOEO 10-Fold Cross-Validation Per-Fold Results')

    # S4: Model Comparison
    doc.add_heading('S4. Model Comparison', level=1)
    
    comp_table = doc.add_table(rows=4, cols=5)
    comp_table.style = 'Table Grid'
    
    comp_headers = ['Model', 'Parameters', 'Mag Acc (LOEO)', 'Azi Acc (LOEO)', 'Model Size']
    for i, h in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = h
    
    models_data = [
        ['VGG16', '138M', '98.68%', '54.93%', '528 MB'],
        ['EfficientNet-B0', '5.3M', '97.53% ± 0.96%', '69.51% ± 5.65%', '20 MB'],
        ['ConvNeXt-Tiny', '28.6M', f"{results['magnitude_accuracy']['mean']:.2f}% ± {results['magnitude_accuracy']['std']:.2f}%", 
         f"{results['azimuth_accuracy']['mean']:.2f}% ± {results['azimuth_accuracy']['std']:.2f}%", '112 MB'] if results else 
        ['ConvNeXt-Tiny', '28.6M', 'N/A', 'N/A', '112 MB']
    ]
    
    for idx, data in enumerate(models_data):
        for j, val in enumerate(data):
            comp_table.rows[idx + 1].cells[j].text = val
    
    doc.add_paragraph()
    doc.add_paragraph('Table S2: Model Comparison Summary')
    
    # S5: Figures
    doc.add_page_break()
    doc.add_heading('S5. Supplementary Figures', level=1)
    
    supp_figures = [
        ('fig4_loeo_boxplot.png', 'Figure S1: LOEO Results Distribution (Box Plot)'),
        ('fig6_fold_heatmap.png', 'Figure S2: Per-Fold Performance Heatmap'),
        ('fig7_sample_distribution.png', 'Figure S3: Sample Distribution per Fold'),
        ('fig8_summary_table.png', 'Figure S4: Results Summary Table'),
    ]
    
    for fig_file, caption in supp_figures:
        fig_path = FIGURES_DIR / fig_file
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(5.5))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    
    # Save
    output_path = OUTPUT_DIR / "ConvNeXt_Supplementary_Materials.docx"
    doc.save(str(output_path))
    print(f"✓ Saved: {output_path}")

if __name__ == "__main__":
    print("Generating Supplementary Materials...")
    create_supplementary()
    print("Done!")
